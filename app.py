"""
AI Career Assistant

Takes your resume and a job description, sends them to the OpenAI API,
and generates a structured career analysis, tailored resume suggestions,
and a cover letter — all saved as local files.
"""

import os
from datetime import datetime

from dotenv import load_dotenv
from openai import OpenAI
from docx import Document


load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def print_header(title: str) -> None:
    print(f"\n{'=' * 40}")
    print(f"  {title}")
    print(f"{'=' * 40}\n")


def load_resume(filepath: str = "resume.txt") -> str:
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        print(f"Error: '{filepath}' not found. Make sure it's in the same directory.")
        exit(1)


def get_job_description() -> str:
    print("Paste the job description below.")
    print("When finished, type END on a new line.\n")
    lines = []
    while True:
        line = input()
        if line.strip().upper() == "END":
            break
        lines.append(line)
    return "\n".join(lines)


SYSTEM_PROMPT = """
You are an experienced IT recruiter, ATS optimization specialist, and career advisor.

Core behavior rules:
1. Prioritize honesty and realism over sounding impressive.
2. Never exaggerate candidate experience or invent qualifications.
3. Write naturally — like a thoughtful, knowledgeable person, not a chatbot.
4. Avoid corporate buzzwords, generic motivational language, and repetitive AI phrasing.
5. Keep writing professional but grounded and human.
6. Vary sentence structure naturally.
7. Tailor tone to context: warm for cover letters, clear and direct for resume tips, honest for match analysis.
8. Keep ATS optimization realistic and truthful.
9. Do not fabricate technologies, certifications, or work experience.
10. Explain skill gaps honestly and constructively.
11. Make cover letters sound authentic and believable — like a real candidate wrote them.
12. Prioritize practical usefulness over polished-sounding output.
"""


def build_prompt(resume_text: str, job_description: str) -> str:
    return f"""
Analyze the candidate's resume against the provided job description.
Return your response in the EXACT format below. Use the section delimiters exactly as shown.

===TAILORED_RESUME===

RESUME MATCH ANALYSIS

1. REALISTIC MATCH SCORE
Give a realistic percentage match with a brief explanation. Be honest — do not inflate the score.

2. STRONGEST MATCHING SKILLS
- skill 1
- skill 2

3. IMPORTANT ATS KEYWORDS
- keyword 1
- keyword 2

4. MISSING OR WEAK AREAS
- missing skill 1
- missing skill 2

5. TAILORED PROFESSIONAL SUMMARY
Write a 4–6 line professional summary tailored to this specific role.
Keep it realistic, human-sounding, and ATS-friendly.
Do not exaggerate or invent experience.

6. TAILORED SKILLS SECTION
Rewrite the skills section to better align with what this role is looking for.
Keep it truthful — only include skills the candidate actually has.

6A. TAILORED EXPERIENCE BULLETS
Rewrite the candidate's experience bullets to better match the job description.
- Keep everything truthful
- Use stronger action verbs
- Improve ATS keyword alignment
- Do not invent tools or technologies not in the original resume

7. RESUME IMPROVEMENT SUGGESTIONS
Give 3–5 specific, practical suggestions to improve the resume for this role.

8. LIKELY RECRUITER CONCERNS
List 3–5 realistic concerns a recruiter might have when reviewing this application.

9. INTERVIEW TOPICS TO PREPARE
- topic 1
- topic 2

===TAILORED_COVER_LETTER===

10. TAILORED COVER LETTER
Write a cover letter tailored to this specific role and candidate.
Requirements:
- Sound human and conversational — not like a template
- Professional but not stiff
- 300–400 words
- Reference specific technologies or responsibilities from the job description
- Be honest about experience level
- Do not invent qualifications or fake enthusiasm
- Structure: short intro, 2 body paragraphs, brief close

===CAREER_ANALYSIS===

Return this section in the following exact format. Do not change the headers or numbering.

==============================
JOB MATCH ANALYSIS
==============================

1. ATS KEYWORDS
List the most important ATS keywords found in the job description.
- keyword 1
- keyword 2

2. TECHNICAL SKILLS REQUIRED
List the key technical skills the job requires.
- skill 1
- skill 2

3. SOFT SKILLS REQUIRED
List the soft skills mentioned or implied in the job description.
- skill 1
- skill 2

4. REALISTIC MATCH SCORE
Format it exactly like this — nothing else:

MATCH SCORE: XX%
Reason: One clear sentence explaining the score honestly.

5. SKILL GAPS
List specific skills or experience the candidate is missing for this role.
- gap 1
- gap 2

6. INTERVIEW TOPICS TO PREPARE
List the most important topics the candidate should prepare for.
- topic 1
- topic 2

7. SHOULD I APPLY?
Give honest, direct advice in 3–5 sentences.
Consider experience level, skill gaps, role fit, and whether applying is worth the candidate's time.
Do not be overly encouraging. Be realistic.


CANDIDATE RESUME

{resume_text}

JOB DESCRIPTION

{job_description}
"""


def call_openai(prompt: str) -> str:
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.7,
    )
    return response.choices[0].message.content


def parse_sections(analysis: str) -> tuple[str, str, str]:
    try:
        resume_section = (
            analysis
            .split("===TAILORED_RESUME===")[1]
            .split("===TAILORED_COVER_LETTER===")[0]
            .strip()
        )
        cover_letter_section = (
            analysis
            .split("===TAILORED_COVER_LETTER===")[1]
            .split("===CAREER_ANALYSIS===")[0]
            .strip()
        )
        career_analysis_section = (
            analysis
            .split("===CAREER_ANALYSIS===")[1]
            .strip()
        )
        return resume_section, cover_letter_section, career_analysis_section

    except (IndexError, Exception) as e:
        print(f"\nWarning: Could not parse sections cleanly ({e}).")
        print("Printing full response instead.\n")
        print(analysis)
        return "", "", ""


def save_outputs(
    career_analysis: str,
    resume_section: str,
    cover_letter: str,
    resume_text: str,
    job_description: str,
) -> None:

    os.makedirs("analyses", exist_ok=True)
    os.makedirs("tailored_resumes", exist_ok=True)
    os.makedirs("cover_letters", exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")

    # Career analysis — saved as .txt with full context
    analysis_path = f"analyses/career_analysis_{timestamp}.txt"

    with open(analysis_path, "w", encoding="utf-8") as f:
        f.write("==============================\n")
        f.write("CANDIDATE RESUME\n")
        f.write("==============================\n\n")
        f.write(resume_text)
        f.write("\n\n==============================\n")
        f.write("JOB DESCRIPTION\n")
        f.write("==============================\n\n")
        f.write(job_description)
        f.write("\n\n==============================\n")
        f.write("AI CAREER ANALYSIS\n")
        f.write("==============================\n")
        f.write("See section 4 below for your MATCH SCORE\n")
        f.write("------------------------------\n\n")
        f.write(career_analysis)

    # Tailored resume — saved as .docx
    resume_path = f"tailored_resumes/tailored_resume_{timestamp}.docx"

    resume_doc = Document()
    resume_doc.add_heading("Tailored Resume Suggestions", level=1)
    resume_doc.add_paragraph(resume_section)
    resume_doc.save(resume_path)

    # Cover letter — saved as .docx
    cover_letter_path = f"cover_letters/cover_letter_{timestamp}.docx"

    cover_doc = Document()
    cover_doc.add_heading("Cover Letter", level=1)
    cover_doc.add_paragraph(cover_letter)
    cover_doc.save(cover_letter_path)

    print_header("FILES GENERATED SUCCESSFULLY")
    print(f"  → {analysis_path}")
    print(f"  → {resume_path}")
    print(f"  → {cover_letter_path}")
    print()


def main():
    print_header("AI CAREER ASSISTANT")

    resume_text = load_resume()
    job_description = get_job_description()

    print("\nAnalyzing your resume against the job description...")
    print("This may take 15–30 seconds depending on length.\n")

    prompt = build_prompt(resume_text, job_description)

    try:
        analysis = call_openai(prompt)
    except Exception as e:
        print(f"\nOpenAI API error: {e}")
        exit(1)

    resume_section, cover_letter_section, career_analysis_section = parse_sections(analysis)

    if career_analysis_section:
        save_outputs(
            career_analysis_section,
            resume_section,
            cover_letter_section,
            resume_text,
            job_description,
        )
    else:
        print("\nNo structured output was saved due to parsing issues.")
        print("Check the output above for the full AI response.")


if __name__ == "__main__":
    main()