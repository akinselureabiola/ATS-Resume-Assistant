from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# ==========================================
# DOCUMENT STYLING
# ==========================================

def style_document(doc):

    sections = doc.sections

    for section in sections:

        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]

    font = style.font

    font.name = "Calibri"
    font.size = Pt(11)


# ==========================================
# NAME HEADER
# ==========================================

def add_name_heading(doc, text):

    paragraph = doc.add_paragraph()

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run(text)

    run.bold = True
    run.font.size = Pt(20)


# ==========================================
# CONTACT LINE
# ==========================================

def add_contact_line(doc, text):

    paragraph = doc.add_paragraph()

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run(text)

    run.font.size = Pt(10)


# ==========================================
# SECTION HEADING
# ==========================================

def add_section_heading(doc, text):

    heading = doc.add_paragraph()

    run = heading.add_run(text)

    run.bold = True
    run.font.size = Pt(14)


# ==========================================
# RESUME GENERATION
# ==========================================

def generate_resume_docx(content, output_path):

    doc = Document()

    style_document(doc)

    paragraphs = content.split("\n")

    first_line = True

    section_titles = [
        "Professional Summary",
        "Skills",
        "Professional Experience",
        "Projects",
        "Technical Project",
        "Education",
        "Certifications",
        "Languages",
        "Profile"
    ]

    for paragraph in paragraphs:

        paragraph = paragraph.strip()

        if not paragraph:
            continue

        # ==========================
        # NAME
        # ==========================

        if first_line:

            add_name_heading(
                doc,
                paragraph
            )

            first_line = False

            continue

        # ==========================
        # CONTACT INFO
        # ==========================

        if (
            "@" in paragraph
            or "linkedin.com" in paragraph
            or "github.com" in paragraph
            or "+" in paragraph
            or "Berlin" in paragraph
            or "Germany" in paragraph
        ):

            add_contact_line(
                doc,
                paragraph
            )

            continue

        # ==========================
        # SECTION HEADINGS
        # ==========================

        if paragraph in section_titles:

            add_section_heading(
                doc,
                paragraph
            )

            continue

        # ==========================
        # BULLET POINTS
        # ==========================

        if (
            paragraph.startswith("-")
            or paragraph.startswith("•")
        ):

            cleaned = (
                paragraph
                .replace("•", "")
                .replace("-", "")
                .strip()
            )

            doc.add_paragraph(
                cleaned,
                style="List Bullet"
            )

            continue

        # ==========================
        # NORMAL TEXT
        # ==========================

        doc.add_paragraph(paragraph)

    doc.save(output_path)

    return output_path


# ==========================================
# COVER LETTER GENERATION
# ==========================================

def generate_cover_letter_docx(content, output_path):

    doc = Document()

    style_document(doc)

    paragraphs = content.split("\n")

    for paragraph in paragraphs:

        paragraph = paragraph.strip()

        if not paragraph:
            continue

        para = doc.add_paragraph()

        run = para.add_run(paragraph)

        run.font.size = Pt(11)

    doc.save(output_path)

    return output_path